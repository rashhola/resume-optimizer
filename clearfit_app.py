import streamlit as st
import requests
from docx import Document
from io import BytesIO
import PyPDF2

# Get Groq API key from secrets
groq_api_key = st.secrets["groq"]["api_key"]

st.title("ClearFit: Tailor Your Resume to Any Job (Llama 3)")

st.write("**Step 1:** Paste the job description below.")
job_desc = st.text_area("Job Description", height=200)

st.write("**Step 2:** Upload your current resume (PDF, DOCX, or TXT).")
resume_file = st.file_uploader("Upload Resume", type=["pdf", "docx", "txt"])

def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    file_type = uploaded_file.type
    if file_type == "application/pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_type == "text/plain":
        return uploaded_file.read().decode("utf-8")
    else:
        return ""

def llama3_generate(prompt):
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {groq_api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "llama3-70b-8192",
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 1024,
        "temperature": 0.7
    }
    response = requests.post(url, headers=headers, json=data)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"].strip()

if st.button("Tailor My Resume"):
    if not job_desc.strip():
        st.error("Please paste the job description.")
    elif not resume_file:
        st.error("Please upload your current resume.")
    else:
        with st.spinner("Tailoring your resume..."):
            user_resume = extract_text_from_file(resume_file)
            if not user_resume.strip():
                st.error("Could not extract text from your resume file.")
            else:
                prompt = (
                    "You are a professional resume writer. Given the following job description and my current resume, "
                    "rewrite my resume so it best matches the job description. Highlight relevant skills, experience, and deliverables. "
                    "Keep the format professional and concise.\n\n"
                    f"Job Description:\n{job_desc}\n\n"
                    f"My Resume:\n{user_resume}\n\n"
                    "Tailored Resume:"
                )
                try:
                    tailored_resume = llama3_generate(prompt)
                except Exception as e:
                    st.error(f"Error: {e}")
                    tailored_resume = None

                if tailored_resume:
                    st.subheader("Tailored Resume")
                    st.text_area("Preview", tailored_resume, height=400)

                    # Download as DOCX
                    doc = Document()
                    for line in tailored_resume.split('\n'):
                        if line.strip() == "":
                            doc.add_paragraph()
                        else:
                            doc.add_paragraph(line)
                    bio = BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    st.download_button(
                        label="Download Tailored Resume as DOCX",
                        data=bio,
                        file_name="Tailored_Resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )